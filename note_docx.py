# coding=utf-8
'''
   Python 读取 word 文本
   Python可以利用python-docx模块处理word文档，处理方式是面向对象的。
   也就是说python-docx模块会把word文档，文档中的段落、文本、字体等都看做对象，对对象进行处理就是对word文档的内容处理。
    1，Document对象，表示一个word文档。
    2，Paragraph对象，表示word文档中的一个段落
    3，Paragraph对象的text属性，表示段落中的文本内容。

    使用样式—— 使用Python读写Office文档:
    https://zhuanlan.zhihu.com/p/23708800?utm_source=tuicool&utm_medium=referral

    Python读写docx文件
    https://blog.csdn.net/g0ose/article/details/64538787
    对象集合
    doc.paragraphs    #段落集合
    doc.tables        #表格集合
    doc.sections      #节  集合
    doc.styles        #样式集合
    doc.inline_shapes #内置图形 等等...
'''
import os
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import win32com
from win32com.client import Dispatch, constants

if __name__ == '__main__':
    file = docx.Document('./Data/Albania2.docx')

    #print("段落数：",str(len(file.paragraphs)))
    # 输出每一段落的内容
    # for para in file.paragraphs:
        # print(para.text)
    #输出段落编号及段落内容
    #for i in range(len(file.paragraphs)):
        #print("第 ",str(i)," 段的内容是：",file.paragraphs[i].text)

def parse_doc(f):
    doc = w.Documents.Open(FileName=f)
    print("段落数：", str(len(doc.paragraphs)))
    for i in range(len(doc.paragraphs)):
        print("第 ", str(i), " 段的内容是：", doc.paragraphs[i].text)
def parse_docx(f):
    d = Document(f)
    print(doc,"段落数：", str(len(d.paragraphs)))
    for i in range(len(d.paragraphs)):
        print("第 ", str(i), " 段的内容是：", d.paragraphs[i].text)
    Force = d.paragraphs[12].text
    print("Enter into force:",Force)
    Promulgation = d.paragraphs[14].text
    print('Promulgation:',d.paragraphs[14].text)

if __name__ == "__main__":
  w = win32com.client.Dispatch('Word.Application')
  # 遍历文件
  PATH = "./BITS" # windows文件路径
  doc_files = os.listdir(PATH)
  for doc in doc_files:
    if os.path.splitext(doc)[1] == '.docx':
      try:
        parse_docx(PATH+'\\'+doc)
      except Exception as e:
        print(e)
    elif os.path.splitext(doc)[1] == '.doc':
      try:
        parse_doc(PATH+'\\'+doc)
      except Exception as e:
        print(e)